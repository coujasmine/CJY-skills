#!/usr/bin/env python3
"""Post-process a Pandoc-generated docx: re-apply body font, add heading numbers,
line numbers, page numbers, and first-line indent. Tables are handled separately
in format_apa_tables.py.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).resolve().parent))
from _lib import load_journal_config


def _force_run_rfonts(run_element, font_name):
    rpr = run_element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)


def apply_body_font(doc, body_spec):
    """Force the body font on every run in non-heading paragraphs.

    Pandoc sometimes leaves runs with no explicit font (so Word falls back to its
    default Calibri). This sweep is defensive — if a run already has the right
    font it's a no-op.
    """
    font_name = body_spec["font_name"]
    font_size_pt = body_spec["font_size_pt"]
    for para in doc.paragraphs:
        style_name = para.style.name if (para.style and para.style.name) else ""
        if style_name.startswith("Heading"):
            continue
        for run in para.runs:
            _force_run_rfonts(run._element, font_name)
            if run.font.size is None:
                run.font.size = Pt(font_size_pt)


def apply_heading_numbering(doc, headings_spec):
    """Prepend '1.', '1.1.', etc. to Heading 1-N paragraphs in document order.

    Skips paragraphs that already start with the correct number.
    Strips any pre-existing leading number that doesn't match (e.g., user wrote
    '## 2 Theory' but the counter says '1' — the script overwrites).
    """
    if headings_spec.get("numbering") != "decimal":
        return
    max_level = headings_spec.get("max_level", 4)
    counters = [0] * max_level

    for para in doc.paragraphs:
        style_name = para.style.name if (para.style and para.style.name) else ""
        m = re.match(r"Heading (\d+)$", style_name)
        if not m:
            continue
        level = int(m.group(1))
        if level < 1 or level > max_level:
            continue
        counters[level - 1] += 1
        for i in range(level, max_level):
            counters[i] = 0
        number = ".".join(str(c) for c in counters[:level])

        current = para.text
        # already-numbered case (idempotent)
        if re.match(rf"^{re.escape(number)}[.\s]", current):
            continue

        # strip any pre-existing leading "N", "N.M", "N.M.O" with optional dot+space
        stripped = re.sub(r"^\s*\d+(\.\d+)*\.?\s+", "", current)
        prefix = f"{number}. "

        if stripped != current and para.runs:
            # We had a stale prefix; wipe and rewrite the first run, blank the rest
            para.runs[0].text = prefix + stripped
            for r in para.runs[1:]:
                r.text = ""
        elif para.runs:
            # Clean prepend onto the first run
            para.runs[0].text = prefix + para.runs[0].text
        else:
            para.add_run(prefix)


def apply_line_numbers(doc, page_spec):
    if not page_spec.get("line_numbers"):
        return
    restart = page_spec.get("line_number_restart", "continuous")
    for section in doc.sections:
        sectPr = section._sectPr
        existing = sectPr.find(qn("w:lnNumType"))
        if existing is not None:
            sectPr.remove(existing)
        lnNumType = OxmlElement("w:lnNumType")
        lnNumType.set(qn("w:countBy"), "1")
        lnNumType.set(qn("w:start"), "1")
        lnNumType.set(qn("w:restart"), restart)
        sectPr.append(lnNumType)


_ALIGN_FOR_POSITION = {
    "bottom_center": ("footer", WD_ALIGN_PARAGRAPH.CENTER),
    "bottom_right": ("footer", WD_ALIGN_PARAGRAPH.RIGHT),
    "bottom_left": ("footer", WD_ALIGN_PARAGRAPH.LEFT),
    "top_center": ("header", WD_ALIGN_PARAGRAPH.CENTER),
    "top_right": ("header", WD_ALIGN_PARAGRAPH.RIGHT),
    "top_left": ("header", WD_ALIGN_PARAGRAPH.LEFT),
}


def apply_page_numbers(doc, page_spec):
    if not page_spec.get("page_numbers"):
        return
    where, align = _ALIGN_FOR_POSITION.get(
        page_spec.get("page_number_position", "bottom_center"),
        ("footer", WD_ALIGN_PARAGRAPH.CENTER),
    )
    for section in doc.sections:
        target = section.footer if where == "footer" else section.header
        para = target.paragraphs[0] if target.paragraphs else target.add_paragraph()
        for r in list(para.runs):
            r._element.getparent().remove(r._element)
        para.alignment = align

        run = para.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._element.append(fldChar1)
        run._element.append(instr)
        run._element.append(fldChar2)


def apply_first_line_indent(doc, body_spec):
    indent_cm = body_spec.get("first_line_indent_cm", 0)
    if not indent_cm:
        return
    indent = Cm(indent_cm)
    for para in doc.paragraphs:
        style_name = para.style.name if (para.style and para.style.name) else ""
        if style_name.startswith("Heading") or not para.text.strip():
            continue
        para.paragraph_format.first_line_indent = indent


def main() -> int:
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("docx")
    parser.add_argument("journal")
    parser.add_argument("--no-heading-numbering", action="store_true",
                        help="Skip auto-numbering of headings (use when the source MD has manual numbers)")
    args = parser.parse_args()

    docx_path = Path(args.docx)
    config = load_journal_config(args.journal)

    doc = Document(str(docx_path))
    apply_body_font(doc, config["body"])
    if not args.no_heading_numbering:
        apply_heading_numbering(doc, config.get("headings", {}))
    apply_line_numbers(doc, config["page"])
    apply_page_numbers(doc, config["page"])
    apply_first_line_indent(doc, config["body"])
    doc.save(str(docx_path))
    print(f"Post-processed: {docx_path}"
          + (" (manual heading numbering preserved)" if args.no_heading_numbering else ""))
    return 0


if __name__ == "__main__":
    sys.exit(main())
