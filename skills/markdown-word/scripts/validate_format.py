#!/usr/bin/env python3
"""Check a docx against a journal YAML spec and emit a markdown compliance report.

Exit code = number of failed checks (0 means full PASS).
"""
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).resolve().parent))
from _lib import load_journal_config


def _approx_eq(a, b, eps=0.05):
    if a is None or b is None:
        return a == b
    return abs(a - b) < eps


def check_page(doc, page_spec):
    out = []
    section = doc.sections[0]

    for side, expected_cm in (
        ("top", page_spec["margin_top_cm"]),
        ("bottom", page_spec["margin_bottom_cm"]),
        ("left", page_spec["margin_left_cm"]),
        ("right", page_spec["margin_right_cm"]),
    ):
        actual = getattr(section, f"{side}_margin")
        actual_cm = actual.cm if actual else None
        ok = _approx_eq(actual_cm, expected_cm, eps=0.05)
        out.append({
            "check": f"page margin ({side})",
            "expected": f"{expected_cm:.2f} cm",
            "actual": f"{actual_cm:.2f} cm" if actual_cm is not None else "unset",
            "pass": ok,
        })

    expected_ln = page_spec.get("line_numbers", False)
    actual_ln = section._sectPr.find(qn("w:lnNumType")) is not None
    out.append({
        "check": "line numbers enabled",
        "expected": str(expected_ln),
        "actual": str(actual_ln),
        "pass": expected_ln == actual_ln,
    })

    if page_spec.get("page_numbers"):
        footer_text = "".join(p.text for p in section.footer.paragraphs).strip()
        has_field = section.footer._element.find(".//" + qn("w:instrText")) is not None
        out.append({
            "check": "page-number field present",
            "expected": "PAGE field",
            "actual": "yes" if has_field else f"no (footer text: {footer_text!r})",
            "pass": has_field,
        })

    return out


def check_body(doc, body_spec):
    out = []
    normal = doc.styles["Normal"]

    expected_font = body_spec["font_name"]
    actual_font = normal.font.name
    out.append({
        "check": "Normal style font",
        "expected": expected_font,
        "actual": actual_font or "unset",
        "pass": actual_font == expected_font,
    })

    expected_size = body_spec["font_size_pt"]
    actual_size = normal.font.size.pt if normal.font.size else None
    out.append({
        "check": "Normal style font size",
        "expected": f"{expected_size} pt",
        "actual": f"{actual_size} pt" if actual_size else "unset",
        "pass": _approx_eq(actual_size, expected_size, eps=0.5),
    })

    expected_ls = body_spec["line_spacing"]
    actual_ls = normal.paragraph_format.line_spacing
    out.append({
        "check": "Normal style line spacing",
        "expected": f"{expected_ls}",
        "actual": f"{actual_ls}",
        "pass": _approx_eq(actual_ls, expected_ls, eps=0.05),
    })

    return out


def check_headings(doc, headings_spec):
    out = []
    if headings_spec.get("numbering") != "decimal":
        return out
    import re
    found = 0
    numbered = 0
    for para in doc.paragraphs:
        name = para.style.name if (para.style and para.style.name) else ""
        if not re.match(r"Heading [1-4]$", name):
            continue
        found += 1
        if re.match(r"^\s*\d+(\.\d+)*\.?\s+\S", para.text):
            numbered += 1
    if found == 0:
        out.append({
            "check": "heading numbering",
            "expected": "decimal",
            "actual": "no Heading 1-4 paragraphs found",
            "pass": True,  # vacuously true
        })
    else:
        out.append({
            "check": f"heading numbering ({found} headings)",
            "expected": "all numbered",
            "actual": f"{numbered}/{found} numbered",
            "pass": numbered == found,
        })
    return out


def check_tables(doc, tables_spec):
    out = []
    n = len(doc.tables)
    out.append({
        "check": "table count",
        "expected": ">= 0",
        "actual": str(n),
        "pass": True,
    })
    if n == 0 or tables_spec.get("style") != "apa":
        return out

    apa_ok = 0
    for table in doc.tables:
        rows = list(table.rows)
        if not rows:
            continue
        first_cell = rows[0].cells[0]
        tcPr = first_cell._tc.find(qn("w:tcPr"))
        if tcPr is None:
            continue
        tcBorders = tcPr.find(qn("w:tcBorders"))
        if tcBorders is None:
            continue
        top = tcBorders.find(qn("w:top"))
        left = tcBorders.find(qn("w:left"))
        right = tcBorders.find(qn("w:right"))
        top_ok = top is not None and top.get(qn("w:val")) == "single"
        verticals_off = (
            (left is None or left.get(qn("w:val")) == "nil")
            and (right is None or right.get(qn("w:val")) == "nil")
        )
        if top_ok and verticals_off:
            apa_ok += 1

    out.append({
        "check": "APA three-line styling (header row, no verticals)",
        "expected": f"{n}/{n} tables",
        "actual": f"{apa_ok}/{n} tables",
        "pass": apa_ok == n,
    })
    return out


def render_report(sections):
    lines = ["# Format compliance report", ""]
    overall_pass = True
    fail_count = 0
    for section_name, rows in sections.items():
        lines.append(f"## {section_name}")
        lines.append("")
        lines.append("| Check | Expected | Actual | Result |")
        lines.append("|---|---|---|---|")
        for r in rows:
            mark = "PASS" if r["pass"] else "FAIL"
            if not r["pass"]:
                overall_pass = False
                fail_count += 1
            lines.append(
                f"| {r['check']} | {r['expected']} | {r['actual']} | {mark} |"
            )
        lines.append("")
    lines.insert(1, f"**Overall: {'PASS' if overall_pass else f'FAIL ({fail_count} issue(s))'}**")
    lines.insert(2, "")
    return "\n".join(lines), fail_count


def main() -> int:
    if len(sys.argv) != 3:
        print("usage: validate_format.py <docx> <journal-key>", file=sys.stderr)
        return 2
    docx_path = Path(sys.argv[1])
    journal_key = sys.argv[2]
    config = load_journal_config(journal_key)
    doc = Document(str(docx_path))

    sections = {
        "Page setup": check_page(doc, config["page"]),
        "Body typography": check_body(doc, config["body"]),
        "Headings": check_headings(doc, config.get("headings", {})),
        "Tables": check_tables(doc, config["tables"]),
    }
    report, fail_count = render_report(sections)

    report_path = docx_path.with_name(docx_path.stem + "-report.md")
    report_path.write_text(report, encoding="utf-8")
    print(report)
    print(f"\nReport saved: {report_path}")
    return fail_count


if __name__ == "__main__":
    sys.exit(main())
