#!/usr/bin/env python3
"""Apply APA three-line styling to every table in a docx:
  - Top, header-bottom, and bottom borders only (no verticals, no inner horizontals)
  - Bold first row
  - Specified font/size on cells
  - Repeat first row across page breaks
  - Prevent rows from splitting across page breaks
  - Re-number captions ('Table 1', 'Table 2', ...) in document order
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).resolve().parent))
from _lib import load_journal_config


def _set_cell_borders(cell, top_pt=None, bottom_pt=None,
                      left_pt=None, right_pt=None):
    """Set top/bottom/left/right borders on a single cell.
    Pass `None` (or 0) for an edge to mark it as 'no border'.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    tcBorders = OxmlElement("w:tcBorders")
    for edge, pt in (("top", top_pt), ("left", left_pt),
                     ("bottom", bottom_pt), ("right", right_pt)):
        elem = OxmlElement(f"w:{edge}")
        if not pt:
            elem.set(qn("w:val"), "nil")
        else:
            elem.set(qn("w:val"), "single")
            # OOXML border size unit is eighths of a point
            elem.set(qn("w:sz"), str(max(2, int(pt * 8))))
            elem.set(qn("w:space"), "0")
            elem.set(qn("w:color"), "000000")
        tcBorders.append(elem)
    tcPr.append(tcBorders)


def apply_three_line_borders(table, top_pt, header_bottom_pt, bottom_pt):
    rows = list(table.rows)
    if not rows:
        return
    n = len(rows)
    for r_idx, row in enumerate(rows):
        for cell in row.cells:
            top = top_pt if r_idx == 0 else None
            bot = None
            if r_idx == 0 and n > 1:
                bot = header_bottom_pt
            if r_idx == n - 1:
                bot = bottom_pt
            _set_cell_borders(cell, top_pt=top, bottom_pt=bot,
                              left_pt=None, right_pt=None)


def clear_table_inner_borders(table):
    """Remove tblBorders at the table level so cell-level borders win.

    Pandoc applies the 'Table Grid' or 'Table' style which sets a grid
    via tblBorders/inside* edges. We zero those out so our cell borders show.
    """
    tblPr = table._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        return
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "nil")
        tblBorders.append(e)
    tblPr.append(tblBorders)


def bold_header_row(table):
    if not table.rows:
        return
    for cell in table.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True


def set_table_font(table, font_name, font_size_pt):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size_pt)
                    rpr = run._element.get_or_add_rPr()
                    rfonts = rpr.find(qn("w:rFonts"))
                    if rfonts is None:
                        rfonts = OxmlElement("w:rFonts")
                        rpr.insert(0, rfonts)
                    rfonts.set(qn("w:ascii"), font_name)
                    rfonts.set(qn("w:hAnsi"), font_name)
                    rfonts.set(qn("w:cs"), font_name)
                    rfonts.set(qn("w:eastAsia"), font_name)


def repeat_header_across_pages(table):
    if not table.rows:
        return
    trPr = table.rows[0]._tr.get_or_add_trPr()
    if trPr.find(qn("w:tblHeader")) is None:
        trPr.append(OxmlElement("w:tblHeader"))


def prevent_row_split(table):
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        if trPr.find(qn("w:cantSplit")) is None:
            trPr.append(OxmlElement("w:cantSplit"))


_CAPTION_RE = re.compile(r"^\s*Table\s*(\d+)?\s*[:.]?\s*", re.IGNORECASE)


def _paragraph_text(p_elem) -> str:
    parts = []
    for t in p_elem.iter(qn("w:t")):
        parts.append(t.text or "")
    return "".join(parts)


def _paragraph_style_name(p_elem, styles_part) -> str:
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is None:
        return ""
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        return ""
    style_id = pStyle.get(qn("w:val")) or ""
    if not styles_part:
        return style_id
    # try to resolve human name
    for style in styles_part.element.findall(qn("w:style")):
        if style.get(qn("w:styleId")) == style_id:
            name_el = style.find(qn("w:name"))
            if name_el is not None:
                return name_el.get(qn("w:val")) or style_id
    return style_id


def _is_caption_candidate(text: str, style_name: str) -> bool:
    if "caption" in (style_name or "").lower():
        return True
    return bool(_CAPTION_RE.match(text or ""))


def _rewrite_caption_paragraph(p_elem, prefix, number, sep, title,
                               prefix_bold, title_italic, font_size_pt):
    """Replace all runs in p_elem with formatted caption runs."""
    # remove existing runs
    for r in p_elem.findall(qn("w:r")):
        p_elem.remove(r)

    def make_run(text, *, bold=False, italic=False, line_break=False):
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        if bold:
            rPr.append(OxmlElement("w:b"))
        if italic:
            rPr.append(OxmlElement("w:i"))
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(font_size_pt * 2)))
        rPr.append(sz)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), str(int(font_size_pt * 2)))
        rPr.append(szCs)
        r.append(rPr)
        if line_break:
            r.append(OxmlElement("w:br"))
        else:
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = text
            r.append(t)
        return r

    p_elem.append(make_run(f"{prefix} {number}", bold=prefix_bold))
    if sep == "\n":
        p_elem.append(make_run("", line_break=True))
    elif sep:
        p_elem.append(make_run(sep))
    if title:
        p_elem.append(make_run(title, italic=title_italic))


def renumber_table_captions(doc, tables_spec):
    """Walk body children in order. For each <w:tbl>, look at the previous
    sibling — if it's a caption candidate, rewrite it to 'Table N <sep> <title>'.
    """
    prefix = tables_spec.get("caption_prefix", "Table")
    sep = tables_spec.get("caption_separator", "\n")
    prefix_bold = tables_spec.get("caption_prefix_bold", False)
    title_italic = tables_spec.get("caption_title_italic", True)
    font_size_pt = tables_spec.get("caption_font_size_pt", 12)

    styles_part = doc.part.styles_part if hasattr(doc.part, "styles_part") else None
    body = doc.element.body
    counter = 0
    skipped_no_caption = []

    for elem in list(body.iterchildren()):
        if not elem.tag.endswith("}tbl"):
            continue
        counter += 1
        prev = elem.getprevious()
        if prev is None or not prev.tag.endswith("}p"):
            skipped_no_caption.append(counter)
            continue
        text = _paragraph_text(prev)
        style_name = _paragraph_style_name(prev, styles_part)
        if not _is_caption_candidate(text, style_name):
            skipped_no_caption.append(counter)
            continue
        title = _CAPTION_RE.sub("", text, count=1).strip()
        _rewrite_caption_paragraph(
            prev, prefix=prefix, number=counter, sep=sep, title=title,
            prefix_bold=prefix_bold, title_italic=title_italic,
            font_size_pt=font_size_pt,
        )
    return counter, skipped_no_caption


def main() -> int:
    if len(sys.argv) != 3:
        print("usage: format_apa_tables.py <docx> <journal-key>", file=sys.stderr)
        return 2
    docx_path = Path(sys.argv[1])
    journal_key = sys.argv[2]
    config = load_journal_config(journal_key)
    tables_spec = config["tables"]

    doc = Document(str(docx_path))
    if not doc.tables:
        print("No tables found.")
        doc.save(str(docx_path))
        return 0

    style = tables_spec.get("style", "apa")
    for table in doc.tables:
        if style == "apa":
            clear_table_inner_borders(table)
            apply_three_line_borders(
                table,
                tables_spec["border_top_pt"],
                tables_spec["border_header_bottom_pt"],
                tables_spec["border_bottom_pt"],
            )
        if tables_spec.get("header_bold", True):
            bold_header_row(table)
        set_table_font(
            table,
            tables_spec.get("font_name", "Times New Roman"),
            tables_spec.get("font_size_pt", 11),
        )
        if tables_spec.get("repeat_header_across_pages"):
            repeat_header_across_pages(table)
        if not tables_spec.get("allow_row_split", False):
            prevent_row_split(table)

    n, skipped = renumber_table_captions(doc, tables_spec)
    doc.save(str(docx_path))
    print(f"Formatted {n} table(s) in {docx_path}")
    if skipped:
        print(f"  Tables without a recognized caption (numbers): {skipped}")
        print("  → Add a paragraph like 'Table: Title' immediately before the table.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
